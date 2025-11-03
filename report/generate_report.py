#!/usr/bin/env python3
"""
Generate Summer Training Report for NIVI AI Project
Following Dr. Akhilesh Das Gupta Institute of Professional Studies guidelines
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ['sz', 'val', 'color']:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    
    tcPr.append(tcBorders)

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def create_cover_page(doc):
    """Create the cover page"""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SUMMER TRAINING REPORT 1\n(ES-361)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True
    
    # Add spacing
    for _ in range(2):
        doc.add_paragraph()
    
    # Report title
    report_title = doc.add_paragraph()
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = report_title.add_run('NIVI AI: Development of an Intelligent\nAI-Powered Conversational Assistant')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    run.italic = True
    run.bold = True
    
    # Add spacing
    for _ in range(2):
        doc.add_paragraph()
    
    # Submitted in partial fulfillment
    fulfillment = doc.add_paragraph()
    fulfillment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fulfillment.add_run('Submitted in partial fulfillment of the requirements\nfor the award of the degree of')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.italic = True
    
    # Add spacing
    doc.add_paragraph()
    
    # Bachelor of Technology
    btech = doc.add_paragraph()
    btech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = btech.add_run('Bachelor of Technology\nDepartment of Computer Science & Engineering')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    run.bold = True
    
    # Add spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Submitted by
    submitted_by = doc.add_paragraph()
    submitted_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = submitted_by.add_run('Submitted by:\n[Student Name]\n[Enrollment Number]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # Add spacing
    for _ in range(2):
        doc.add_paragraph()
    
    # Institute name
    institute = doc.add_paragraph()
    institute.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = institute.add_run('Dr. Akhilesh Das Gupta Institute of Professional Studies\n(Formerly ADGITM)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    run.bold = True
    
    # Address
    address = doc.add_paragraph()
    address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = address.add_run('FC-26, SHASTRI PARK, NEW DELHI')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    
    # Add spacing
    doc.add_paragraph()
    
    # Affiliated to
    affiliated = doc.add_paragraph()
    affiliated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = affiliated.add_run('Affiliated to')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    
    # University
    university = doc.add_paragraph()
    university.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = university.add_run('GURU GOBIND SINGH INDRAPRASTHA UNIVERSITY')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    run.bold = True
    
    # University address
    univ_addr = doc.add_paragraph()
    univ_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = univ_addr.add_run('Sector - 16C Dwarka, Delhi - 110075, India')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    
    # Year
    doc.add_paragraph()
    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run('2023-27')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True

def create_declaration(doc):
    """Create declaration page"""
    add_page_break(doc)
    
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('DECLARATION')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    
    declaration_text = doc.add_paragraph()
    declaration_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    declaration_text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = declaration_text.add_run(
        'This is to certify that the material embodied in this Summer Training Report-1 titled '
        '"NIVI AI: Development of an Intelligent AI-Powered Conversational Assistant" being submitted '
        'in the partial fulfillment of the requirements for the award of the degree of Bachelor of '
        'Technology in Computer Science & Engineering is based on my original work. It is further '
        'certified that this work has not been submitted in full or in part to this university or any '
        'other university for the award of any other degree or diploma. My indebtedness to other works '
        'has been duly acknowledged at the relevant places.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    # Add spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Student signature
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = signature.add_run('[Student Name]\n[Enrollment Number]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def create_acknowledgement(doc):
    """Create acknowledgement page"""
    add_page_break(doc)
    
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('ACKNOWLEDGEMENT')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    
    ack_text = doc.add_paragraph()
    ack_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ack_text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = ack_text.add_run(
        'I would like to express my sincere gratitude to my training organization for providing me with '
        'the opportunity to work on this project. Their guidance, support, and invaluable insights '
        'throughout the training period have significantly enhanced my technical skills and professional '
        'development.\n\n'
        'I am deeply thankful to Dr. Akhilesh Das Gupta Institute of Professional Studies for organizing '
        'this summer training program and providing the academic framework that enabled me to undertake '
        'this meaningful project. Special thanks to my faculty mentors and the Training and Placement '
        'Cell for their continuous support.\n\n'
        'I would also like to acknowledge the open-source community and the developers of the various '
        'technologies and frameworks used in this project, particularly React, Vite, and Tailwind CSS, '
        'which made the development process efficient and enjoyable.\n\n'
        'Finally, I extend my heartfelt thanks to my family and friends for their unwavering support '
        'and encouragement throughout this journey.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def create_abstract(doc):
    """Create abstract page"""
    add_page_break(doc)
    
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('ABSTRACT')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    
    abstract_text = doc.add_paragraph()
    abstract_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = abstract_text.add_run(
        'This report presents a comprehensive overview of the summer training program undertaken focused '
        'on Full-Stack Web Development and Artificial Intelligence Integration. The training resulted in '
        'the development of NIVI AI, an intelligent conversational assistant that leverages modern web '
        'technologies and AI capabilities to provide users with an enhanced interactive experience.\n\n'
        'The project encompasses the complete development lifecycle, from initial planning and design to '
        'implementation and deployment. NIVI AI is built using React.js for the frontend, providing a '
        'responsive and intuitive user interface, while integrating advanced AI capabilities for natural '
        'language processing and conversation management.\n\n'
        'Key features implemented include real-time chat functionality, multi-conversation management, '
        'persistent data storage, voice interaction capabilities, analytics dashboard, and comprehensive '
        'user authentication system. The application follows modern software development practices including '
        'component-based architecture, state management, API integration, and responsive design principles.\n\n'
        'The training covered essential concepts in web development including HTML5, CSS3, JavaScript ES6+, '
        'React.js framework, RESTful API design, database management, authentication systems, and deployment '
        'strategies. Special emphasis was placed on creating a scalable, maintainable, and user-friendly '
        'application architecture.\n\n'
        'This report details the technical implementation, challenges encountered, solutions developed, '
        'and the valuable learning outcomes achieved during the four-week intensive training period.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def create_table_of_contents(doc):
    """Create table of contents"""
    add_page_break(doc)
    
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('TABLE OF CONTENTS')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    
    # TOC items
    toc_items = [
        ('Declaration', 'i'),
        ('Acknowledgement', 'ii'),
        ('Abstract', 'iii'),
        ('List of Figures', 'iv'),
        ('List of Tables', 'v'),
        ('Chapter 1: Introduction of the Summer Training Course', '1'),
        ('    1.1: About the Summer Training Course', '1'),
        ('    1.2: Profile of the Training Organization', '2'),
        ('Chapter 2: Introduction of the Training', '4'),
        ('    2.1: About the Training', '4'),
        ('    2.2: Objectives of the Training', '5'),
        ('    2.3: Roles and Responsibilities', '6'),
        ('Chapter 3: Problem Statement', '7'),
        ('    3.1: Software Requirement Specifications', '7'),
        ('        3.1.1: Functional Requirements', '8'),
        ('        3.1.2: Non-functional Requirements', '9'),
        ('    3.2: Feasibility Study', '10'),
        ('    3.3: Tools / Technologies / Platform Used', '11'),
        ('    3.4: System Architecture and Diagrams', '13'),
        ('Chapter 4: Project Activities', '16'),
        ('    4.1: Task Description', '16'),
        ('    4.2: Tools / Technologies / Platform Used', '18'),
        ('    4.3: Technical Application', '20'),
        ('    4.4: Challenges Faced', '24'),
        ('Chapter 5: Learning and Development', '26'),
        ('Chapter 6: Summary and Conclusion', '29'),
        ('Chapter 7: Suggestions for Improvement', '30'),
        ('Bibliography', '31'),
    ]
    
    for item, page_no in toc_items:
        toc_entry = doc.add_paragraph()
        toc_entry.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        toc_entry.paragraph_format.left_indent = Inches(0.5) if item.startswith('    ') else Inches(0)
        
        run1 = toc_entry.add_run(item.strip())
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        
        # Add dots
        dots = '.' * (80 - len(item))
        run2 = toc_entry.add_run(' ' + dots + ' ')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        
        run3 = toc_entry.add_run(page_no)
        run3.font.name = 'Times New Roman'
        run3.font.size = Pt(12)

def create_list_of_figures(doc):
    """Create list of figures"""
    add_page_break(doc)
    
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('LIST OF FIGURES')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Figure No.'
    hdr_cells[1].text = 'Figure Title'
    hdr_cells[2].text = 'Page No.'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Figures
    figures = [
        ('3.1', 'System Architecture Diagram', '13'),
        ('3.2', 'Use Case Diagram', '14'),
        ('3.3', 'Data Flow Diagram', '15'),
        ('4.1', 'NIVI AI Main Interface', '17'),
        ('4.2', 'Chat Conversation Interface', '19'),
        ('4.3', 'Component Architecture', '21'),
        ('4.4', 'Code Structure - React Components', '22'),
        ('4.5', 'Authentication Flow', '23'),
    ]
    
    for fig_no, title, page in figures:
        row_cells = table.add_row().cells
        row_cells[0].text = fig_no
        row_cells[1].text = title
        row_cells[2].text = page
        
        for i, cell in enumerate(row_cells):
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(12)
            if i != 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_list_of_tables(doc):
    """Create list of tables"""
    add_page_break(doc)
    
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('LIST OF TABLES')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Table No.'
    hdr_cells[1].text = 'Table Title'
    hdr_cells[2].text = 'Page No.'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tables
    tables = [
        ('3.1', 'Functional Requirements', '8'),
        ('3.2', 'Non-functional Requirements', '9'),
        ('3.3', 'Technology Stack', '12'),
        ('4.1', 'Development Tasks and Timeline', '16'),
        ('4.2', 'Key Features Implementation', '18'),
    ]
    
    for tbl_no, title, page in tables:
        row_cells = table.add_row().cells
        row_cells[0].text = tbl_no
        row_cells[1].text = title
        row_cells[2].text = page
        
        for i, cell in enumerate(row_cells):
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(12)
            if i != 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    """Main function to generate the report"""
    print("Generating Summer Training Report...")
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set margins (1.5" left, 1" right, top, bottom)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
    
    # Create all pages
    create_cover_page(doc)
    create_declaration(doc)
    create_acknowledgement(doc)
    create_abstract(doc)
    create_table_of_contents(doc)
    create_list_of_figures(doc)
    create_list_of_tables(doc)
    
    # Save document
    output_path = '/home/runner/work/NIVI/NIVI/report/Summer_Training_Report_Part1.docx'
    doc.save(output_path)
    print(f"Part 1 of report saved to: {output_path}")
    
    return output_path

if __name__ == '__main__':
    main()
