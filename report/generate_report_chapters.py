#!/usr/bin/env python3
"""
Generate Summer Training Report Chapters for NIVI AI Project
Part 2: All Chapters with Content
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def add_chapter_heading(doc, chapter_title):
    """Add a chapter heading"""
    heading = doc.add_paragraph()
    run = heading.add_run(chapter_title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

def add_main_heading(doc, heading_text):
    """Add a main heading"""
    heading = doc.add_paragraph()
    run = heading.add_run(heading_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

def add_sub_heading(doc, heading_text):
    """Add a sub heading"""
    heading = doc.add_paragraph()
    run = heading.add_run(heading_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph(doc, text, justify=True):
    """Add a justified paragraph"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return para

def add_bullet_point(doc, text):
    """Add a bullet point"""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def create_chapter_1(doc):
    """Create Chapter 1: Introduction of the Summer Training Course"""
    add_page_break(doc)
    add_chapter_heading(doc, 'CHAPTER 1: INTRODUCTION OF THE SUMMER TRAINING COURSE')
    
    add_main_heading(doc, '1.1 About the Summer Training Course')
    
    add_paragraph(doc, 
        'The Summer Training Course is a mandatory component of the Bachelor of Technology program at '
        'Dr. Akhilesh Das Gupta Institute of Professional Studies, designed to bridge the gap between '
        'theoretical knowledge and practical industry experience. This training program, conducted after '
        'the completion of the 4th semester, provides students with hands-on exposure to real-world '
        'software development practices and emerging technologies.')
    
    add_paragraph(doc,
        'The primary objective of this training is to enable students to apply their academic learning '
        'in a practical setting, thereby enhancing their technical skills, problem-solving abilities, and '
        'industry readiness. The training duration ranges from 4 to 6 weeks, during which students are '
        'expected to work on substantial projects that demonstrate their capabilities in software design, '
        'development, and deployment.')
    
    add_paragraph(doc,
        'As part of the course requirements, students must develop a comprehensive project that showcases '
        'their understanding of software engineering principles, modern development frameworks, and best '
        'practices in the industry. The training culminates in a detailed technical report and presentation '
        'before a faculty evaluation committee.')
    
    add_main_heading(doc, '1.2 Profile of the Training Organization')
    
    add_paragraph(doc,
        'The summer training was undertaken as an independent learning and development initiative, focusing '
        'on modern web development technologies and artificial intelligence integration. The training '
        'environment simulated real-world software development scenarios, incorporating industry-standard '
        'tools, frameworks, and methodologies.')
    
    add_paragraph(doc,
        'The training program emphasized:')
    
    add_bullet_point(doc, 'Full-stack web application development using modern JavaScript frameworks')
    add_bullet_point(doc, 'Integration of AI and machine learning capabilities in web applications')
    add_bullet_point(doc, 'Responsive and user-centric interface design')
    add_bullet_point(doc, 'RESTful API development and integration')
    add_bullet_point(doc, 'Database management and data persistence')
    add_bullet_point(doc, 'Authentication and authorization systems')
    add_bullet_point(doc, 'Deployment and hosting strategies')
    
    add_paragraph(doc,
        'The training provided exposure to cutting-edge technologies including React.js, Node.js, modern '
        'CSS frameworks, and cloud-based AI services. This comprehensive approach ensured a holistic '
        'understanding of contemporary software development practices.')

def create_chapter_2(doc):
    """Create Chapter 2: Introduction of the Training"""
    add_page_break(doc)
    add_chapter_heading(doc, 'CHAPTER 2: INTRODUCTION OF THE TRAINING')
    
    add_main_heading(doc, '2.1 About the Training')
    
    add_paragraph(doc,
        'The summer training focused on developing NIVI AI, an intelligent conversational assistant that '
        'demonstrates the practical application of modern web technologies and artificial intelligence. '
        'The training spanned four weeks of intensive learning and development, covering the entire '
        'software development lifecycle from conceptualization to deployment.')
    
    add_paragraph(doc,
        'The training methodology combined self-directed learning with practical implementation, requiring '
        'extensive research, experimentation, and iterative development. Each week focused on specific '
        'aspects of the application, building progressively from basic functionality to advanced features.')
    
    add_paragraph(doc,
        'Week-wise breakdown:')
    
    add_bullet_point(doc, 'Week 1: Project planning, technology stack selection, and environment setup')
    add_bullet_point(doc, 'Week 2: Core functionality development including chat interface and AI integration')
    add_bullet_point(doc, 'Week 3: Advanced features implementation including authentication and data persistence')
    add_bullet_point(doc, 'Week 4: Testing, optimization, documentation, and deployment')
    
    add_main_heading(doc, '2.2 Objectives of the Training')
    
    add_paragraph(doc,
        'The primary objectives of this summer training were:')
    
    add_bullet_point(doc, 
        'To develop proficiency in modern web development frameworks, particularly React.js and its ecosystem')
    add_bullet_point(doc,
        'To understand and implement AI integration in web applications for natural language processing')
    add_bullet_point(doc,
        'To gain hands-on experience with component-based architecture and state management')
    add_bullet_point(doc,
        'To implement responsive design principles for cross-device compatibility')
    add_bullet_point(doc,
        'To develop a complete full-stack application from scratch')
    add_bullet_point(doc,
        'To understand and implement user authentication and authorization mechanisms')
    add_bullet_point(doc,
        'To learn database integration and data persistence strategies')
    add_bullet_point(doc,
        'To implement modern UI/UX design principles for enhanced user experience')
    add_bullet_point(doc,
        'To develop problem-solving and debugging skills through practical challenges')
    add_bullet_point(doc,
        'To understand deployment and hosting processes for web applications')
    
    add_main_heading(doc, '2.3 Roles and Responsibilities')
    
    add_paragraph(doc,
        'As the sole developer on this training project, the responsibilities encompassed all aspects of '
        'software development:')
    
    add_sub_heading(doc, 'Project Planning and Design:')
    add_bullet_point(doc, 'Defining project scope and objectives')
    add_bullet_point(doc, 'Creating technical specifications and requirements documentation')
    add_bullet_point(doc, 'Designing system architecture and component structure')
    add_bullet_point(doc, 'Creating wireframes and user interface mockups')
    
    add_sub_heading(doc, 'Development:')
    add_bullet_point(doc, 'Implementing frontend components using React.js')
    add_bullet_point(doc, 'Integrating AI capabilities for intelligent conversations')
    add_bullet_point(doc, 'Developing state management and data flow architecture')
    add_bullet_point(doc, 'Creating responsive and accessible user interfaces')
    add_bullet_point(doc, 'Implementing authentication and authorization systems')
    
    add_sub_heading(doc, 'Testing and Quality Assurance:')
    add_bullet_point(doc, 'Conducting functional testing of all features')
    add_bullet_point(doc, 'Performing cross-browser compatibility testing')
    add_bullet_point(doc, 'Debugging and resolving issues')
    add_bullet_point(doc, 'Optimizing performance and responsiveness')
    
    add_sub_heading(doc, 'Documentation:')
    add_bullet_point(doc, 'Maintaining comprehensive code documentation')
    add_bullet_point(doc, 'Creating user guides and technical documentation')
    add_bullet_point(doc, 'Documenting API endpoints and integration points')
    add_bullet_point(doc, 'Preparing this detailed training report')

def create_chapter_3(doc):
    """Create Chapter 3: Problem Statement"""
    add_page_break(doc)
    add_chapter_heading(doc, 'CHAPTER 3: PROBLEM STATEMENT')
    
    add_paragraph(doc,
        'The project aimed to develop NIVI AI, an intelligent conversational assistant that provides users '
        'with an intuitive interface for natural language interactions. The challenge was to create a '
        'feature-rich, scalable, and user-friendly application that combines modern web technologies with '
        'AI capabilities while maintaining high performance and excellent user experience.')
    
    add_main_heading(doc, '3.1 Software Requirement Specifications')
    
    add_sub_heading(doc, '3.1.1 Functional Requirements')
    
    add_paragraph(doc,
        'The following table outlines the functional requirements of NIVI AI:')
    
    doc.add_paragraph()
    
    # Create functional requirements table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Requirement ID'
    hdr_cells[1].text = 'Description'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Requirements
    requirements = [
        ('FR-1', 'Real-time conversational interface with AI-powered responses'),
        ('FR-2', 'Multi-conversation management with persistent storage'),
        ('FR-3', 'User authentication and authorization system'),
        ('FR-4', 'Chat history with search and filter capabilities'),
        ('FR-5', 'Voice input and output functionality'),
        ('FR-6', 'File attachment and sharing capabilities'),
        ('FR-7', 'Export chat conversations in multiple formats'),
        ('FR-8', 'Analytics dashboard for usage tracking'),
        ('FR-9', 'Customizable settings and preferences'),
        ('FR-10', 'Responsive design for all device types'),
        ('FR-11', 'Quick prompt suggestions for common tasks'),
        ('FR-12', 'Message reactions and bookmarking features'),
    ]
    
    for req_id, desc in requirements:
        row_cells = table.add_row().cells
        row_cells[0].text = req_id
        row_cells[1].text = desc
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    add_sub_heading(doc, '3.1.2 Non-functional Requirements')
    
    add_paragraph(doc,
        'Non-functional requirements define the quality attributes of the system:')
    
    doc.add_paragraph()
    
    # Create non-functional requirements table
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    
    # Header
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Requirement ID'
    hdr_cells[1].text = 'Description'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Requirements
    nf_requirements = [
        ('NFR-1', 'Performance: Response time under 2 seconds for normal queries'),
        ('NFR-2', 'Scalability: Support for concurrent users without degradation'),
        ('NFR-3', 'Usability: Intuitive interface with minimal learning curve'),
        ('NFR-4', 'Reliability: 99.5% uptime with automatic error recovery'),
        ('NFR-5', 'Security: Encrypted data transmission and secure authentication'),
        ('NFR-6', 'Maintainability: Modular code structure for easy updates'),
        ('NFR-7', 'Compatibility: Cross-browser support for modern browsers'),
        ('NFR-8', 'Accessibility: WCAG 2.1 Level AA compliance'),
    ]
    
    for req_id, desc in nf_requirements:
        row_cells = table2.add_row().cells
        row_cells[0].text = req_id
        row_cells[1].text = desc
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    add_main_heading(doc, '3.2 Feasibility Study')
    
    add_sub_heading(doc, 'Technical Feasibility:')
    add_paragraph(doc,
        'The project is technically feasible with the selected technology stack. React.js provides a robust '
        'framework for building interactive user interfaces, while modern AI APIs offer reliable natural '
        'language processing capabilities. The availability of comprehensive documentation and community '
        'support ensures successful implementation.')
    
    add_sub_heading(doc, 'Economic Feasibility:')
    add_paragraph(doc,
        'The project utilizes primarily open-source technologies, minimizing development costs. The AI '
        'integration uses a freemium model with generous free tiers suitable for development and testing. '
        'Hosting can be accomplished using free or low-cost cloud platforms.')
    
    add_sub_heading(doc, 'Operational Feasibility:')
    add_paragraph(doc,
        'The application is designed for ease of use, requiring minimal training. The intuitive interface '
        'and familiar chat paradigm ensure high user acceptance. Maintenance is simplified through modular '
        'architecture and comprehensive documentation.')
    
    add_main_heading(doc, '3.3 Tools / Technologies / Platform Used')
    
    add_paragraph(doc,
        'The following table summarizes the complete technology stack:')
    
    doc.add_paragraph()
    
    # Create technology table
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    
    # Header
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Purpose'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Technologies
    technologies = [
        ('Frontend Framework', 'React.js 18.2', 'UI component development'),
        ('Build Tool', 'Vite 7.1', 'Fast development and building'),
        ('Styling', 'Tailwind CSS 3.3', 'Responsive utility-first styling'),
        ('State Management', 'React Hooks', 'Application state management'),
        ('Routing', 'React Router 7.9', 'Client-side routing'),
        ('Icons', 'Lucide React', 'Modern icon library'),
        ('AI Integration', 'Advanced AI API', 'Natural language processing'),
        ('Storage', 'LocalStorage & MongoDB', 'Data persistence'),
        ('Authentication', 'JWT Tokens', 'Secure user authentication'),
        ('PDF Generation', 'jsPDF & html2canvas', 'Export functionality'),
        ('ID Generation', 'nanoid', 'Unique identifier generation'),
        ('Development', 'VS Code', 'Integrated development environment'),
        ('Version Control', 'Git & GitHub', 'Source code management'),
    ]
    
    for category, tech, purpose in technologies:
        row_cells = table3.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = tech
        row_cells[2].text = purpose
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    add_main_heading(doc, '3.4 System Architecture and Diagrams')
    
    add_paragraph(doc,
        'NIVI AI follows a component-based architecture pattern, separating concerns between presentation, '
        'business logic, and data management layers. The application uses a client-side rendering approach '
        'with React.js, communicating with AI services through RESTful APIs.')
    
    add_sub_heading(doc, 'Architecture Components:')
    add_bullet_point(doc, 'Presentation Layer: React components for UI rendering')
    add_bullet_point(doc, 'Business Logic Layer: Custom hooks and utility functions')
    add_bullet_point(doc, 'Data Layer: State management and API integration')
    add_bullet_point(doc, 'Storage Layer: LocalStorage and database connectivity')
    
    add_paragraph(doc,
        '[Figure 3.1: System Architecture Diagram would be inserted here]')
    
    add_paragraph(doc,
        '[Figure 3.2: Use Case Diagram would be inserted here]')
    
    add_paragraph(doc,
        '[Figure 3.3: Data Flow Diagram would be inserted here]')

def create_chapter_4(doc):
    """Create Chapter 4: Project Activities"""
    add_page_break(doc)
    add_chapter_heading(doc, 'CHAPTER 4: PROJECT ACTIVITIES')
    
    add_main_heading(doc, '4.1 Task Description')
    
    add_paragraph(doc,
        'The development of NIVI AI was divided into several key phases, each focusing on specific '
        'functionalities and features. The following table outlines the major tasks and their timeline:')
    
    doc.add_paragraph()
    
    # Create task table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Tasks'
    hdr_cells[2].text = 'Duration'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tasks
    tasks = [
        ('Planning', 'Requirements gathering, technology research, architecture design', '3 days'),
        ('Setup', 'Environment configuration, project initialization, dependency installation', '2 days'),
        ('Core Development', 'Chat interface, AI integration, message handling', '5 days'),
        ('Feature Development', 'Multi-chat, authentication, voice mode, analytics', '7 days'),
        ('UI/UX Enhancement', 'Styling, responsive design, animations, accessibility', '4 days'),
        ('Testing', 'Functional testing, bug fixes, performance optimization', '3 days'),
        ('Documentation', 'Code documentation, user guides, technical documentation', '2 days'),
        ('Deployment', 'Build optimization, hosting setup, final testing', '2 days'),
    ]
    
    for phase, task, duration in tasks:
        row_cells = table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = task
        row_cells[2].text = duration
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    add_paragraph(doc,
        '[Figure 4.1: NIVI AI Main Interface screenshot would be inserted here]')
    
    add_main_heading(doc, '4.2 Tools / Technologies / Platform Used')
    
    add_paragraph(doc,
        'The development process leveraged various tools and technologies for different aspects:')
    
    add_sub_heading(doc, 'Development Tools:')
    add_bullet_point(doc, 'Visual Studio Code: Primary code editor with extensions for React development')
    add_bullet_point(doc, 'Chrome DevTools: Debugging and performance analysis')
    add_bullet_point(doc, 'React Developer Tools: Component inspection and profiling')
    add_bullet_point(doc, 'Git: Version control and code management')
    add_bullet_point(doc, 'npm: Package management and dependency handling')
    
    add_sub_heading(doc, 'Key Libraries and Frameworks:')
    add_bullet_point(doc, 'React.js: Frontend framework for component-based UI development')
    add_bullet_point(doc, 'Vite: Next-generation build tool for fast development')
    add_bullet_point(doc, 'Tailwind CSS: Utility-first CSS framework for rapid styling')
    add_bullet_point(doc, 'React Router: Declarative routing for React applications')
    add_bullet_point(doc, 'Lucide React: Icon library for consistent iconography')
    
    add_paragraph(doc,
        'The following table summarizes the key features and their implementation technologies:')
    
    doc.add_paragraph()
    
    # Feature implementation table
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Technology Used'
    hdr_cells[2].text = 'Implementation Details'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    features = [
        ('Chat Interface', 'React Hooks', 'useState and useEffect for state management'),
        ('AI Integration', 'Fetch API', 'RESTful API calls to AI service'),
        ('Authentication', 'JWT, Context API', 'Token-based auth with React Context'),
        ('Data Persistence', 'LocalStorage, MongoDB', 'Dual storage strategy'),
        ('Routing', 'React Router', 'Dynamic routing for multiple chats'),
        ('Export', 'jsPDF, html2canvas', 'PDF generation from chat content'),
        ('Analytics', 'Custom utility functions', 'Usage tracking and statistics'),
        ('Voice Mode', 'Web Speech API', 'Browser-native speech recognition'),
    ]
    
    for feature, tech, details in features:
        row_cells = table2.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = tech
        row_cells[2].text = details
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    add_paragraph(doc,
        '[Figure 4.2: Chat Conversation Interface screenshot would be inserted here]')
    
    add_main_heading(doc, '4.3 Technical Application')
    
    add_paragraph(doc,
        'The technical implementation of NIVI AI demonstrates various software engineering principles and '
        'modern development practices:')
    
    add_sub_heading(doc, 'Component Architecture:')
    add_paragraph(doc,
        'The application follows a modular component structure, with clear separation of concerns. Main '
        'components include:')
    
    add_bullet_point(doc, 'App.jsx: Root component managing application state and routing')
    add_bullet_point(doc, 'ChatArea.jsx: Main chat interface with message display and input')
    add_bullet_point(doc, 'Sidebar.jsx: Navigation and chat history management')
    add_bullet_point(doc, 'ChatMessage.jsx: Individual message rendering component')
    add_bullet_point(doc, 'AuthPage.jsx: Authentication interface for login/registration')
    add_bullet_point(doc, 'SettingsModal.jsx: User preferences and configuration')
    add_bullet_point(doc, 'AnalyticsModal.jsx: Usage statistics and insights')
    
    add_paragraph(doc,
        '[Figure 4.3: Component Architecture diagram would be inserted here]')
    
    add_sub_heading(doc, 'State Management:')
    add_paragraph(doc,
        'React Hooks (useState, useEffect, useContext) manage application state efficiently. Key state '
        'variables include chat history, current messages, user authentication status, and UI preferences. '
        'The Context API provides global state for authentication and theme management.')
    
    add_sub_heading(doc, 'API Integration:')
    add_paragraph(doc,
        'The application integrates with AI services through RESTful APIs. A sample implementation for '
        'sending messages to the AI:')
    
    # Add code block representation
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_run = code_para.add_run(
        'const sendMessage = async (message) => {\n'
        '  const response = await fetch(API_URL, {\n'
        '    method: "POST",\n'
        '    headers: { "Content-Type": "application/json" },\n'
        '    body: JSON.stringify({ message })\n'
        '  });\n'
        '  const data = await response.json();\n'
        '  return data.response;\n'
        '};'
    )
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    add_paragraph(doc,
        '[Figure 4.4: Code Structure - React Components would be inserted here]')
    
    add_sub_heading(doc, 'Data Persistence:')
    add_paragraph(doc,
        'The application implements a dual storage strategy: LocalStorage for unauthenticated users and '
        'MongoDB for authenticated users. This ensures data persistence while providing flexibility for '
        'different user types.')
    
    add_sub_heading(doc, 'Authentication Flow:')
    add_paragraph(doc,
        'JWT-based authentication provides secure user management. The authentication flow includes '
        'registration, login, token storage, and automatic token refresh. Protected routes ensure that '
        'certain features are only accessible to authenticated users.')
    
    add_paragraph(doc,
        '[Figure 4.5: Authentication Flow diagram would be inserted here]')
    
    add_main_heading(doc, '4.4 Challenges Faced')
    
    add_paragraph(doc,
        'During the development process, several challenges were encountered and resolved:')
    
    add_sub_heading(doc, 'Challenge 1: State Management Complexity')
    add_paragraph(doc,
        'Managing multiple interdependent state variables across components proved challenging, especially '
        'with chat history synchronization. Solution: Implemented a centralized state management approach '
        'using React Context API and custom hooks to ensure consistent state across components.')
    
    add_sub_heading(doc, 'Challenge 2: API Rate Limiting')
    add_paragraph(doc,
        'AI API rate limits could cause service interruptions during heavy usage. Solution: Implemented '
        'response caching, request queuing, and retry logic with exponential backoff to handle rate limits '
        'gracefully.')
    
    add_sub_heading(doc, 'Challenge 3: Real-time Updates')
    add_paragraph(doc,
        'Ensuring smooth real-time updates without performance degradation required optimization. Solution: '
        'Used React\'s reconciliation algorithm efficiently, implemented virtual scrolling for long chat '
        'histories, and optimized re-rendering with React.memo and useCallback.')
    
    add_sub_heading(doc, 'Challenge 4: Cross-browser Compatibility')
    add_paragraph(doc,
        'Certain features like voice recognition had varying support across browsers. Solution: Implemented '
        'progressive enhancement with feature detection, providing fallback options for unsupported browsers.')
    
    add_sub_heading(doc, 'Challenge 5: Data Synchronization')
    add_paragraph(doc,
        'Synchronizing data between LocalStorage and MongoDB for users who later authenticate was complex. '
        'Solution: Developed a migration utility that seamlessly transfers local data to the database upon '
        'user authentication.')

def create_chapter_5(doc):
    """Create Chapter 5: Learning and Development"""
    add_page_break(doc)
    add_chapter_heading(doc, 'CHAPTER 5: LEARNING AND DEVELOPMENT')
    
    add_paragraph(doc,
        'The summer training program provided extensive learning opportunities across multiple domains of '
        'software development. The hands-on experience with NIVI AI development significantly enhanced both '
        'technical and soft skills.')
    
    add_main_heading(doc, 'Technical Skills Acquired')
    
    add_sub_heading(doc, 'Frontend Development:')
    add_bullet_point(doc, 'Mastered React.js framework including hooks, context, and advanced patterns')
    add_bullet_point(doc, 'Gained proficiency in modern JavaScript (ES6+) features and best practices')
    add_bullet_point(doc, 'Learned responsive design principles using Tailwind CSS')
    add_bullet_point(doc, 'Understood component lifecycle and optimization techniques')
    add_bullet_point(doc, 'Developed skills in creating accessible and user-friendly interfaces')
    
    add_sub_heading(doc, 'State Management:')
    add_bullet_point(doc, 'Implemented complex state management using React Hooks')
    add_bullet_point(doc, 'Utilized Context API for global state management')
    add_bullet_point(doc, 'Developed custom hooks for reusable logic')
    add_bullet_point(doc, 'Managed asynchronous state updates effectively')
    
    add_sub_heading(doc, 'API Integration:')
    add_bullet_point(doc, 'Gained experience in integrating third-party APIs')
    add_bullet_point(doc, 'Learned RESTful API principles and best practices')
    add_bullet_point(doc, 'Implemented error handling and retry mechanisms')
    add_bullet_point(doc, 'Developed understanding of AI service integration')
    
    add_sub_heading(doc, 'Database and Storage:')
    add_bullet_point(doc, 'Learned data persistence strategies using LocalStorage')
    add_bullet_point(doc, 'Gained experience with MongoDB for cloud storage')
    add_bullet_point(doc, 'Implemented data migration and synchronization logic')
    add_bullet_point(doc, 'Understood data modeling for chat applications')
    
    add_sub_heading(doc, 'Authentication and Security:')
    add_bullet_point(doc, 'Implemented JWT-based authentication system')
    add_bullet_point(doc, 'Learned secure password handling and token management')
    add_bullet_point(doc, 'Developed understanding of OAuth and authorization flows')
    add_bullet_point(doc, 'Implemented protected routes and role-based access')
    
    add_sub_heading(doc, 'Build Tools and Development Workflow:')
    add_bullet_point(doc, 'Mastered Vite build tool for fast development')
    add_bullet_point(doc, 'Learned module bundling and code splitting')
    add_bullet_point(doc, 'Gained experience with hot module replacement')
    add_bullet_point(doc, 'Understood production build optimization')
    
    add_main_heading(doc, 'Soft Skills Development')
    
    add_sub_heading(doc, 'Problem-Solving:')
    add_paragraph(doc,
        'Encountered numerous technical challenges that required analytical thinking and creative solutions. '
        'Developed systematic debugging approaches and learned to break complex problems into manageable '
        'components.')
    
    add_sub_heading(doc, 'Self-Learning:')
    add_paragraph(doc,
        'The training required extensive self-directed learning through documentation, online tutorials, '
        'and community resources. Developed ability to quickly learn new technologies and frameworks '
        'independently.')
    
    add_sub_heading(doc, 'Time Management:')
    add_paragraph(doc,
        'Managing the four-week timeline with multiple deliverables enhanced project planning and time '
        'management skills. Learned to prioritize tasks and maintain development velocity.')
    
    add_sub_heading(doc, 'Documentation:')
    add_paragraph(doc,
        'Created comprehensive code documentation, user guides, and this technical report. Developed '
        'technical writing skills and learned importance of clear documentation.')
    
    add_sub_heading(doc, 'Attention to Detail:')
    add_paragraph(doc,
        'Ensuring consistent UI/UX, handling edge cases, and maintaining code quality required meticulous '
        'attention to detail. Learned importance of thorough testing and code review.')
    
    add_main_heading(doc, 'Industry-Relevant Competencies')
    
    add_bullet_point(doc, 'Understanding of software development lifecycle (SDLC)')
    add_bullet_point(doc, 'Experience with version control using Git')
    add_bullet_point(doc, 'Knowledge of agile development methodologies')
    add_bullet_point(doc, 'Familiarity with modern web development stack')
    add_bullet_point(doc, 'Understanding of UX/UI design principles')
    add_bullet_point(doc, 'Experience with testing and debugging strategies')
    add_bullet_point(doc, 'Knowledge of deployment and hosting processes')
    add_bullet_point(doc, 'Understanding of code optimization and performance')
    
    add_main_heading(doc, 'Certifications and Additional Learning')
    
    add_paragraph(doc,
        'During the training period, supplementary learning included:')
    
    add_bullet_point(doc, 'Completed online courses on React.js advanced concepts')
    add_bullet_point(doc, 'Studied documentation for AI API integration')
    add_bullet_point(doc, 'Reviewed best practices for web application security')
    add_bullet_point(doc, 'Explored accessibility guidelines (WCAG 2.1)')
    add_bullet_point(doc, 'Studied modern web performance optimization techniques')
    
    add_main_heading(doc, 'Key Takeaways')
    
    add_paragraph(doc,
        'The most valuable lessons learned during this training include:')
    
    add_bullet_point(doc, 
        'Importance of planning and architecture before implementation')
    add_bullet_point(doc,
        'Value of writing clean, maintainable, and well-documented code')
    add_bullet_point(doc,
        'Significance of user experience in application success')
    add_bullet_point(doc,
        'Need for continuous learning in rapidly evolving technology landscape')
    add_bullet_point(doc,
        'Importance of testing and handling edge cases')
    add_bullet_point(doc,
        'Value of iterative development and continuous improvement')

def create_chapter_6(doc):
    """Create Chapter 6: Summary and Conclusion"""
    add_page_break(doc)
    add_chapter_heading(doc, 'CHAPTER 6: SUMMARY AND CONCLUSION')
    
    add_paragraph(doc,
        'The summer training program culminated in the successful development of NIVI AI, a comprehensive '
        'intelligent conversational assistant that demonstrates practical application of modern web '
        'technologies and artificial intelligence integration.')
    
    add_main_heading(doc, 'Project Summary')
    
    add_paragraph(doc,
        'NIVI AI represents a full-featured conversational AI platform built entirely using contemporary '
        'web development technologies. The application successfully implements:')
    
    add_bullet_point(doc, 'Real-time chat interface with AI-powered natural language understanding')
    add_bullet_point(doc, 'Multi-conversation management with persistent data storage')
    add_bullet_point(doc, 'Comprehensive user authentication and authorization system')
    add_bullet_point(doc, 'Voice interaction capabilities for hands-free operation')
    add_bullet_point(doc, 'Analytics dashboard for usage insights and statistics')
    add_bullet_point(doc, 'Export functionality for saving conversations')
    add_bullet_point(doc, 'Responsive design ensuring cross-device compatibility')
    add_bullet_point(doc, 'Accessible interface following modern UX/UI principles')
    
    add_main_heading(doc, 'Achievement of Objectives')
    
    add_paragraph(doc,
        'All primary objectives outlined at the beginning of the training were successfully achieved:')
    
    add_bullet_point(doc,
        'Developed comprehensive understanding of React.js and modern JavaScript ecosystem')
    add_bullet_point(doc,
        'Successfully integrated AI capabilities for intelligent conversation management')
    add_bullet_point(doc,
        'Implemented complete authentication and authorization system')
    add_bullet_point(doc,
        'Created responsive, accessible, and user-friendly interface')
    add_bullet_point(doc,
        'Developed problem-solving skills through overcoming technical challenges')
    add_bullet_point(doc,
        'Gained practical experience in full-stack application development')
    
    add_main_heading(doc, 'Technical Accomplishments')
    
    add_paragraph(doc,
        'From a technical perspective, the project demonstrates proficiency in:')
    
    add_bullet_point(doc, 'Component-based architecture and React development patterns')
    add_bullet_point(doc, 'State management using hooks and Context API')
    add_bullet_point(doc, 'RESTful API integration and asynchronous programming')
    add_bullet_point(doc, 'Database integration and data persistence strategies')
    add_bullet_point(doc, 'Modern CSS frameworks and responsive design')
    add_bullet_point(doc, 'Build tools and development workflow optimization')
    add_bullet_point(doc, 'Security best practices and authentication implementation')
    
    add_main_heading(doc, 'Learning Outcomes')
    
    add_paragraph(doc,
        'The training provided invaluable learning experiences that extend beyond technical skills:')
    
    add_paragraph(doc,
        'Technical Growth: Gained proficiency in modern web development stack, understanding of software '
        'architecture principles, and hands-on experience with industry-standard tools and practices.')
    
    add_paragraph(doc,
        'Professional Development: Enhanced problem-solving abilities, improved time management skills, '
        'developed self-learning capabilities, and gained appreciation for documentation and code quality.')
    
    add_paragraph(doc,
        'Industry Readiness: The project provided practical experience that bridges the gap between academic '
        'learning and industry requirements, preparing for professional software development roles.')
    
    add_main_heading(doc, 'Conclusion')
    
    add_paragraph(doc,
        'The four-week summer training program was highly successful in achieving its educational and '
        'developmental objectives. The NIVI AI project not only demonstrates technical competency but also '
        'reflects understanding of user needs, software design principles, and modern development practices.')
    
    add_paragraph(doc,
        'The experience gained during this training has significantly enhanced readiness for professional '
        'software development roles. The practical exposure to real-world development challenges, combined '
        'with the opportunity to implement a complete application from conception to deployment, has provided '
        'invaluable insights into the software development lifecycle.')
    
    add_paragraph(doc,
        'The training has instilled confidence in the ability to learn new technologies, solve complex '
        'problems, and deliver quality software products. The skills and knowledge acquired will serve as '
        'a strong foundation for future academic and professional endeavors in the field of computer science '
        'and software engineering.')

def create_chapter_7(doc):
    """Create Chapter 7: Suggestions for Improvement"""
    add_page_break(doc)
    add_chapter_heading(doc, 'CHAPTER 7: SUGGESTIONS FOR IMPROVEMENT')
    
    add_paragraph(doc,
        'Based on the experience and learning outcomes from this training, the following suggestions are '
        'proposed for future enhancements to both the project and the training program:')
    
    add_main_heading(doc, 'Technical Enhancements for NIVI AI')
    
    add_sub_heading(doc, '1. Advanced AI Capabilities:')
    add_bullet_point(doc, 'Implement multi-modal AI support for handling images and documents')
    add_bullet_point(doc, 'Add context-aware responses using conversation memory')
    add_bullet_point(doc, 'Integrate specialized AI models for domain-specific queries')
    add_bullet_point(doc, 'Implement real-time language translation for multilingual support')
    
    add_sub_heading(doc, '2. Enhanced User Experience:')
    add_bullet_point(doc, 'Add customizable themes beyond dark/light modes')
    add_bullet_point(doc, 'Implement rich text formatting in messages')
    add_bullet_point(doc, 'Add emoji and GIF support for more expressive communication')
    add_bullet_point(doc, 'Implement progressive web app (PWA) capabilities for offline access')
    
    add_sub_heading(doc, '3. Collaboration Features:')
    add_bullet_point(doc, 'Add ability to share conversations with other users')
    add_bullet_point(doc, 'Implement collaborative chat sessions')
    add_bullet_point(doc, 'Add team workspaces for organizational use')
    add_bullet_point(doc, 'Implement conversation templates for common use cases')
    
    add_sub_heading(doc, '4. Performance Optimization:')
    add_bullet_point(doc, 'Implement server-side rendering for improved initial load')
    add_bullet_point(doc, 'Add code splitting for larger bundle optimization')
    add_bullet_point(doc, 'Implement service workers for better caching')
    add_bullet_point(doc, 'Optimize database queries and indexing strategies')
    
    add_sub_heading(doc, '5. Advanced Features:')
    add_bullet_point(doc, 'Add integration with popular productivity tools')
    add_bullet_point(doc, 'Implement scheduled messages and reminders')
    add_bullet_point(doc, 'Add voice cloning for personalized responses')
    add_bullet_point(doc, 'Implement conversation summarization for long chats')
    
    add_main_heading(doc, 'Training Program Improvements')
    
    add_sub_heading(doc, '1. Structured Mentorship:')
    add_paragraph(doc,
        'While self-directed learning was valuable, regular mentorship sessions with industry professionals '
        'would provide guided learning and faster resolution of technical challenges.')
    
    add_sub_heading(doc, '2. Collaborative Projects:')
    add_paragraph(doc,
        'Incorporating team-based projects would provide experience in collaboration, code reviews, and '
        'working with version control in team settings.')
    
    add_sub_heading(doc, '3. Industry Exposure:')
    add_paragraph(doc,
        'Organizing sessions with industry experts, technical talks, and exposure to real-world enterprise '
        'applications would enhance understanding of industry practices.')
    
    add_sub_heading(doc, '4. Extended Duration:')
    add_paragraph(doc,
        'Extending the training to 8 weeks would allow for more comprehensive feature development and '
        'exploration of advanced concepts.')
    
    add_sub_heading(doc, '5. Focus Areas:')
    add_paragraph(doc,
        'Additional focus on testing frameworks, CI/CD pipelines, containerization, and cloud deployment '
        'would better prepare students for industry requirements.')
    
    add_main_heading(doc, 'Documentation and Knowledge Sharing')
    
    add_bullet_point(doc, 'Create video tutorials documenting the development process')
    add_bullet_point(doc, 'Develop comprehensive API documentation')
    add_bullet_point(doc, 'Prepare case studies on specific technical challenges')
    add_bullet_point(doc, 'Establish knowledge base for common issues and solutions')
    
    add_main_heading(doc, 'Future Roadmap')
    
    add_paragraph(doc,
        'Looking ahead, the following roadmap is proposed for NIVI AI:')
    
    add_paragraph(doc,
        'Short-term (1-3 months):')
    add_bullet_point(doc, 'Implement automated testing suite')
    add_bullet_point(doc, 'Add multi-language support')
    add_bullet_point(doc, 'Enhance mobile responsiveness')
    add_bullet_point(doc, 'Implement advanced analytics')
    
    add_paragraph(doc,
        'Medium-term (3-6 months):')
    add_bullet_point(doc, 'Develop native mobile applications')
    add_bullet_point(doc, 'Add enterprise features and admin panel')
    add_bullet_point(doc, 'Implement advanced AI models')
    add_bullet_point(doc, 'Add integration marketplace')
    
    add_paragraph(doc,
        'Long-term (6-12 months):')
    add_bullet_point(doc, 'Scale to support millions of users')
    add_bullet_point(doc, 'Develop AI assistant marketplace')
    add_bullet_point(doc, 'Implement federated learning for privacy')
    add_bullet_point(doc, 'Add blockchain for secure data management')

def create_bibliography(doc):
    """Create bibliography"""
    add_page_break(doc)
    add_chapter_heading(doc, 'BIBLIOGRAPHY')
    
    # References in IEEE format
    references = [
        '[1] React Documentation, "React – A JavaScript library for building user interfaces," '
        'React.dev, 2023. [Online]. Available: https://react.dev/',
        
        '[2] MDN Web Docs, "JavaScript," Mozilla Developer Network, 2023. [Online]. '
        'Available: https://developer.mozilla.org/en-US/docs/Web/JavaScript',
        
        '[3] Vite, "Vite | Next Generation Frontend Tooling," Vitejs.dev, 2023. [Online]. '
        'Available: https://vitejs.dev/',
        
        '[4] Tailwind CSS, "Tailwind CSS - Rapidly build modern websites," Tailwindcss.com, 2023. '
        '[Online]. Available: https://tailwindcss.com/',
        
        '[5] React Router, "React Router: Declarative Routing for React," Reactrouter.com, 2023. '
        '[Online]. Available: https://reactrouter.com/',
        
        '[6] A. Banks and E. Porcello, "Learning React: Modern Patterns for Developing React Apps," '
        'O\'Reilly Media, 2nd ed., 2020.',
        
        '[7] F. Chodorow and D. Dirolf, "MongoDB: The Definitive Guide," O\'Reilly Media, 3rd ed., 2019.',
        
        '[8] M. Fowler, "Patterns of Enterprise Application Architecture," Addison-Wesley Professional, 2002.',
        
        '[9] Stack Overflow, "Stack Overflow Developer Survey 2023," StackOverflow, 2023. [Online]. '
        'Available: https://stackoverflow.com/',
        
        '[10] GitHub, "GitHub: Where the world builds software," GitHub.com, 2023. [Online]. '
        'Available: https://github.com/',
    ]
    
    for ref in references:
        ref_para = doc.add_paragraph()
        ref_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        ref_para.paragraph_format.left_indent = Inches(0.5)
        ref_para.paragraph_format.first_line_indent = Inches(-0.5)
        run = ref_para.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def main():
    """Main function to generate all chapters"""
    print("Generating Summer Training Report Chapters...")
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
    
    # Create all chapters
    create_chapter_1(doc)
    create_chapter_2(doc)
    create_chapter_3(doc)
    create_chapter_4(doc)
    create_chapter_5(doc)
    create_chapter_6(doc)
    create_chapter_7(doc)
    create_bibliography(doc)
    
    # Save document
    output_path = '/home/runner/work/NIVI/NIVI/report/Summer_Training_Report_Chapters.docx'
    doc.save(output_path)
    print(f"Chapters saved to: {output_path}")
    
    return output_path

if __name__ == '__main__':
    main()
