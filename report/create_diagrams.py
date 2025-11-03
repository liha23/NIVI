#!/usr/bin/env python3
"""
Generate Complete Summer Training Report WITH IMAGES
This version includes embedded certificate, offer letter, diagrams, and screenshots
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import os

REPORT_DIR = '/home/runner/work/NIVI/NIVI/report'

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def add_chapter_heading(doc, chapter_title):
    """Add a chapter heading"""
    heading = doc.add_paragraph()
    run = heading.add_run(chapter_title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

def add_main_heading(doc, heading_text):
    """Add a main heading"""
    heading = doc.add_paragraph()
    run = heading.add_run(heading_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

def add_sub_heading(doc, heading_text):
    """Add a sub heading"""
    heading = doc.add_paragraph()
    run = heading.add_run(heading_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph(doc, text, justify=True):
    """Add a justified paragraph"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return para

def add_bullet_point(doc, text):
    """Add a bullet point"""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_image_with_caption(doc, image_path, caption, figure_num, width_inches=6.0):
    """Add an image with caption"""
    try:
        doc.add_paragraph()
        
        # Add image
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        
        # Add caption
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(f'Figure {figure_num}: {caption}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        
        doc.add_paragraph()
        return True
    except Exception as e:
        print(f"Error adding image {image_path}: {e}")
        # Add placeholder instead
        placeholder = doc.add_paragraph()
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = placeholder.add_run(f'[Figure {figure_num}: {caption}]')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True
        doc.add_paragraph()
        return False

def create_diagram_system_architecture():
    """Create system architecture diagram"""
    output_path = os.path.join(REPORT_DIR, 'diagrams', 'system_architecture.png')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Create image
    width, height = 1200, 800
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    # Colors
    box_color = (70, 130, 180)  # Steel blue
    text_color = (255, 255, 255)  # White
    border_color = (0, 0, 0)  # Black
    arrow_color = (100, 100, 100)  # Gray
    
    # Draw title
    title_font_size = 32
    draw.text((width//2, 40), 'NIVI AI - System Architecture', fill=border_color, anchor='mm')
    
    # Define boxes (x, y, width, height, label)
    boxes = [
        (100, 150, 300, 100, 'Frontend Layer\n(React.js)\n- Components\n- State Management\n- Routing'),
        (500, 150, 300, 100, 'User Interface\n- Chat Display\n- Input Area\n- Sidebar'),
        (900, 150, 250, 100, 'Visual Elements\n- Animations\n- Icons\n- Themes'),
        (300, 320, 300, 100, 'API Layer\n(RESTful)\n- Authentication\n- Chat Management'),
        (700, 320, 300, 100, 'Storage Layer\n- LocalStorage\n- MongoDB'),
        (400, 490, 400, 100, 'AI Service Layer\n- Natural Language Processing\n- Response Generation'),
    ]
    
    # Draw boxes
    for x, y, w, h, label in boxes:
        # Draw box
        draw.rectangle([x, y, x+w, y+h], fill=box_color, outline=border_color, width=3)
        # Draw text
        lines = label.split('\n')
        line_height = 20
        start_y = y + h//2 - (len(lines) * line_height)//2
        for i, line in enumerate(lines):
            text_y = start_y + i * line_height
            draw.text((x + w//2, text_y), line, fill=text_color, anchor='mm')
    
    # Draw arrows
    arrows = [
        (450, 200, 500, 200),  # Frontend to UI
        (800, 200, 900, 200),  # UI to Visual
        (450, 250, 450, 320),  # Frontend to API
        (650, 370, 700, 370),  # API to Storage
        (600, 420, 600, 490),  # API to AI
    ]
    
    for x1, y1, x2, y2 in arrows:
        draw.line([x1, y1, x2, y2], fill=arrow_color, width=3)
        # Draw arrowhead
        if x2 > x1:  # Right arrow
            draw.polygon([(x2, y2), (x2-10, y2-5), (x2-10, y2+5)], fill=arrow_color)
        elif y2 > y1:  # Down arrow
            draw.polygon([(x2, y2), (x2-5, y2-10), (x2+5, y2-10)], fill=arrow_color)
    
    img.save(output_path)
    print(f"Created system architecture diagram: {output_path}")
    return output_path

def create_diagram_use_case():
    """Create use case diagram"""
    output_path = os.path.join(REPORT_DIR, 'diagrams', 'use_case.png')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Create image
    width, height = 1200, 900
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    # Colors
    actor_color = (100, 149, 237)  # Cornflower blue
    usecase_color = (255, 218, 185)  # Peach
    border_color = (0, 0, 0)
    
    # Draw title
    draw.text((width//2, 40), 'NIVI AI - Use Case Diagram', fill=border_color, anchor='mm')
    
    # Draw actor (stick figure)
    actor_x, actor_y = 150, 400
    draw.ellipse([actor_x-20, actor_y-40, actor_x+20, actor_y], fill=actor_color, outline=border_color, width=2)
    draw.line([actor_x, actor_y, actor_x, actor_y+60], fill=border_color, width=3)
    draw.line([actor_x-30, actor_y+20, actor_x+30, actor_y+20], fill=border_color, width=3)
    draw.line([actor_x, actor_y+60, actor_x-25, actor_y+100], fill=border_color, width=3)
    draw.line([actor_x, actor_y+60, actor_x+25, actor_y+100], fill=border_color, width=3)
    draw.text((actor_x, actor_y+120), 'User', fill=border_color, anchor='mm')
    
    # Draw use cases (ovals)
    use_cases = [
        (400, 150, 'Login/Register'),
        (400, 250, 'Start New Chat'),
        (400, 350, 'Send Message'),
        (400, 450, 'View History'),
        (700, 150, 'Search Chats'),
        (700, 250, 'Export Chat'),
        (700, 350, 'Use Voice Mode'),
        (700, 450, 'View Analytics'),
        (1000, 250, 'Manage Settings'),
        (1000, 350, 'Bookmark Messages'),
    ]
    
    for x, y, label in use_cases:
        # Draw oval
        draw.ellipse([x-100, y-30, x+100, y+30], fill=usecase_color, outline=border_color, width=2)
        draw.text((x, y), label, fill=border_color, anchor='mm')
        
        # Draw line from actor to use case
        draw.line([actor_x+20, actor_y, x-100, y], fill=border_color, width=1)
    
    # Draw system boundary
    draw.rectangle([350, 100, 1050, 500], outline=border_color, width=3)
    draw.text((700, 80), 'NIVI AI System', fill=border_color, anchor='mm')
    
    img.save(output_path)
    print(f"Created use case diagram: {output_path}")
    return output_path

def create_diagram_data_flow():
    """Create data flow diagram"""
    output_path = os.path.join(REPORT_DIR, 'diagrams', 'data_flow.png')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Create image
    width, height = 1400, 600
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    # Colors
    process_color = (135, 206, 250)  # Light sky blue
    storage_color = (144, 238, 144)  # Light green
    border_color = (0, 0, 0)
    
    # Draw title
    draw.text((width//2, 40), 'NIVI AI - Data Flow Diagram', fill=border_color, anchor='mm')
    
    # Define nodes
    nodes = [
        (100, 200, 150, 80, 'User Input', process_color, 'rect'),
        (350, 200, 150, 80, 'Frontend\nValidation', process_color, 'rect'),
        (600, 200, 150, 80, 'API\nRequest', process_color, 'rect'),
        (850, 200, 150, 80, 'AI\nProcessing', process_color, 'rect'),
        (1100, 200, 150, 80, 'Response\nGeneration', process_color, 'rect'),
        (600, 400, 150, 80, 'Data\nStorage', storage_color, 'rect'),
    ]
    
    # Draw nodes
    for x, y, w, h, label, color, shape in nodes:
        if shape == 'rect':
            draw.rectangle([x, y, x+w, y+h], fill=color, outline=border_color, width=3)
        lines = label.split('\n')
        line_height = 20
        start_y = y + h//2 - (len(lines) * line_height)//2
        for i, line in enumerate(lines):
            text_y = start_y + i * line_height
            draw.text((x + w//2, text_y), line, fill=border_color, anchor='mm')
    
    # Draw arrows
    arrows = [
        (250, 240, 350, 240, 'Message'),
        (500, 240, 600, 240, 'Validated'),
        (750, 240, 850, 240, 'API Call'),
        (1000, 240, 1100, 240, 'AI Result'),
        (675, 280, 675, 400, 'Save'),
        (675, 400, 500, 280, 'Retrieve'),
    ]
    
    for x1, y1, x2, y2, *label in arrows:
        draw.line([x1, y1, x2, y2], fill=border_color, width=3)
        # Draw arrowhead
        if x2 > x1:  # Right arrow
            draw.polygon([(x2, y2), (x2-10, y2-5), (x2-10, y2+5)], fill=border_color)
        elif y2 > y1:  # Down arrow
            draw.polygon([(x2, y2), (x2-5, y2-10), (x2+5, y2-10)], fill=border_color)
        elif x1 > x2:  # Left arrow
            draw.polygon([(x2, y2), (x2+10, y2-5), (x2+10, y2+5)], fill=border_color)
        # Add label
        if label:
            mid_x, mid_y = (x1+x2)//2, (y1+y2)//2 - 15
            draw.text((mid_x, mid_y), label[0], fill=border_color, anchor='mm')
    
    img.save(output_path)
    print(f"Created data flow diagram: {output_path}")
    return output_path

def create_diagram_component_architecture():
    """Create component architecture diagram"""
    output_path = os.path.join(REPORT_DIR, 'diagrams', 'component_architecture.png')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Create image
    width, height = 1200, 1000
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    # Colors
    root_color = (255, 182, 193)  # Light pink
    child_color = (173, 216, 230)  # Light blue
    subchild_color = (255, 228, 181)  # Moccasin
    border_color = (0, 0, 0)
    
    # Draw title
    draw.text((width//2, 40), 'NIVI AI - Component Architecture', fill=border_color, anchor='mm')
    
    # Define tree structure
    components = [
        (600, 100, 200, 60, 'App (Root)', root_color),
        (200, 220, 150, 50, 'AuthProvider', child_color),
        (400, 220, 150, 50, 'ThemeProvider', child_color),
        (600, 220, 150, 50, 'Router', child_color),
        (800, 220, 150, 50, 'Sidebar', child_color),
        (1000, 220, 150, 50, 'ChatArea', child_color),
        (100, 340, 120, 40, 'Login', subchild_color),
        (250, 340, 120, 40, 'Register', subchild_color),
        (750, 340, 120, 40, 'ChatHistory', subchild_color),
        (900, 340, 120, 40, 'SearchBar', subchild_color),
        (950, 340, 120, 40, 'Header', subchild_color),
        (950, 420, 120, 40, 'MessageList', subchild_color),
        (950, 500, 120, 40, 'InputArea', subchild_color),
        (850, 580, 100, 35, 'FileUpload', subchild_color),
        (970, 580, 100, 35, 'VoiceInput', subchild_color),
        (1090, 580, 100, 35, 'SendButton', subchild_color),
    ]
    
    # Draw components
    for x, y, w, h, label, color in components:
        draw.rectangle([x-w//2, y-h//2, x+w//2, y+h//2], fill=color, outline=border_color, width=2)
        lines = label.split('\n')
        line_height = 15
        start_y = y - (len(lines) * line_height)//2
        for i, line in enumerate(lines):
            text_y = start_y + i * line_height
            draw.text((x, text_y), line, fill=border_color, anchor='mm')
    
    # Draw connections
    connections = [
        (600, 160, 275, 220),  # App to AuthProvider
        (600, 160, 475, 220),  # App to ThemeProvider
        (600, 160, 675, 220),  # App to Router
        (600, 160, 875, 220),  # App to Sidebar
        (600, 160, 1075, 220),  # App to ChatArea
        (200, 270, 160, 340),  # AuthProvider to Login
        (200, 270, 310, 340),  # AuthProvider to Register
        (875, 270, 810, 340),  # Sidebar to ChatHistory
        (875, 270, 960, 340),  # Sidebar to SearchBar
        (1075, 270, 1010, 340),  # ChatArea to Header
        (1075, 270, 1010, 420),  # ChatArea to MessageList
        (1075, 270, 1010, 500),  # ChatArea to InputArea
        (1010, 535, 900, 580),  # InputArea to FileUpload
        (1010, 535, 1020, 580),  # InputArea to VoiceInput
        (1010, 535, 1140, 580),  # InputArea to SendButton
    ]
    
    for x1, y1, x2, y2 in connections:
        draw.line([x1, y1, x2, y2], fill=border_color, width=2)
    
    img.save(output_path)
    print(f"Created component architecture diagram: {output_path}")
    return output_path

def create_diagram_auth_flow():
    """Create authentication flow diagram"""
    output_path = os.path.join(REPORT_DIR, 'diagrams', 'auth_flow.png')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Create image
    width, height = 800, 1000
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    # Colors
    start_color = (144, 238, 144)  # Light green
    process_color = (135, 206, 250)  # Light sky blue
    decision_color = (255, 255, 153)  # Light yellow
    border_color = (0, 0, 0)
    
    # Draw title
    draw.text((width//2, 40), 'NIVI AI - Authentication Flow', fill=border_color, anchor='mm')
    
    # Define flow steps
    steps = [
        (400, 100, 100, 50, 'START', start_color, 'oval'),
        (400, 200, 150, 60, 'User Access\nApplication', process_color, 'rect'),
        (400, 320, 120, 80, 'Authenticated?', decision_color, 'diamond'),
        (250, 450, 140, 60, 'Show\nLogin/Register', process_color, 'rect'),
        (550, 450, 140, 60, 'Load User\nData', process_color, 'rect'),
        (250, 570, 140, 60, 'Enter\nCredentials', process_color, 'rect'),
        (250, 690, 140, 60, 'Validate\nCredentials', process_color, 'rect'),
        (250, 810, 140, 60, 'Generate\nJWT Token', process_color, 'rect'),
        (550, 690, 140, 60, 'Show Main\nApplication', process_color, 'rect'),
    ]
    
    # Draw steps
    for x, y, w, h, label, color, shape in steps:
        if shape == 'oval':
            draw.ellipse([x-w//2, y-h//2, x+w//2, y+h//2], fill=color, outline=border_color, width=3)
        elif shape == 'rect':
            draw.rectangle([x-w//2, y-h//2, x+w//2, y+h//2], fill=color, outline=border_color, width=3)
        elif shape == 'diamond':
            points = [(x, y-h//2), (x+w//2, y), (x, y+h//2), (x-w//2, y)]
            draw.polygon(points, fill=color, outline=border_color, width=3)
        
        lines = label.split('\n')
        line_height = 18
        start_y = y - (len(lines) * line_height)//2
        for i, line in enumerate(lines):
            text_y = start_y + i * line_height
            draw.text((x, text_y), line, fill=border_color, anchor='mm')
    
    # Draw arrows with labels
    arrows = [
        (400, 150, 400, 200, ''),
        (400, 260, 400, 280, ''),
        (340, 340, 250, 450, 'No'),
        (460, 340, 550, 450, 'Yes'),
        (250, 510, 250, 570, ''),
        (250, 630, 250, 690, ''),
        (250, 750, 250, 810, ''),
        (320, 810, 550, 750, ''),
        (550, 750, 550, 690, ''),
        (550, 510, 550, 690, ''),
    ]
    
    for x1, y1, x2, y2, label in arrows:
        draw.line([x1, y1, x2, y2], fill=border_color, width=3)
        # Draw arrowhead
        if y2 > y1 and x1 == x2:  # Down arrow
            draw.polygon([(x2, y2), (x2-5, y2-10), (x2+5, y2-10)], fill=border_color)
        elif x2 > x1:  # Right arrow
            draw.polygon([(x2, y2), (x2-10, y2-5), (x2-10, y2+5)], fill=border_color)
        elif x2 < x1:  # Left arrow
            draw.polygon([(x2, y2), (x2+10, y2-5), (x2+10, y2+5)], fill=border_color)
        
        if label:
            mid_x, mid_y = (x1+x2)//2, (y1+y2)//2 - 15
            draw.text((mid_x, mid_y), label, fill=border_color, anchor='mm')
    
    img.save(output_path)
    print(f"Created authentication flow diagram: {output_path}")
    return output_path

def convert_pdf_to_image(pdf_path):
    """Convert first page of PDF to image"""
    try:
        from pdf2image import convert_from_path
        output_path = pdf_path.replace('.pdf', '.png')
        images = convert_from_path(pdf_path, first_page=1, last_page=1)
        if images:
            images[0].save(output_path, 'PNG')
            print(f"Converted PDF to image: {output_path}")
            return output_path
    except Exception as e:
        print(f"Error converting PDF: {e}")
    return None

print("Creating all diagrams...")
create_diagram_system_architecture()
create_diagram_use_case()
create_diagram_data_flow()
create_diagram_component_architecture()
create_diagram_auth_flow()
print("All diagrams created successfully!")
